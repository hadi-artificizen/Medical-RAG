import os, shutil, tempfile, re
from datetime import datetime, timedelta
from typing import List
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials  
from pydantic import BaseModel
from dotenv import load_dotenv
from langchain_community.document_loaders import PyPDFLoader, CSVLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from langchain_openai import ChatOpenAI
from langchain.schema import HumanMessage, SystemMessage
from docx import Document as DocxDocument
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import psycopg2
from fastapi.responses import StreamingResponse


from jose import JWTError, jwt
import bcrypt

load_dotenv()


JWT_SECRET_KEY = "medicalchatbotsecretkey"
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 0.01

security = HTTPBearer()

PERSIST_DIR = "./chroma_db"

chat_llm = ChatOpenAI(
    model="openai/gpt-oss-20b",
    temperature=0.5,
    base_url="https://api.groq.com/openai/v1"
)

SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_USER = "imabdul.hadi1234@gmail.com"
SMTP_PASSWORD = "gqty jnji ufjq rgrh"
EMAIL_TIMEOUT = 30

DB_URL = os.getenv("DB_URL")
conn = psycopg2.connect(DB_URL)
cur = conn.cursor()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)


def hash_password(password: str) -> str:
    
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

def verify_password(plain_password: str, hashed_password: str) -> bool:
    return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))

def create_access_token(data: dict, expires_delta: timedelta = None):
  
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)
    return encoded_jwt

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    
    token = credentials.credentials
    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Could not validate credentials",
                headers={"WWW-Authenticate": "Bearer"},
            )
        return email
    except JWTError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Could not validate credentials",
            headers={"WWW-Authenticate": "Bearer"},
        )

class RegisterRequest(BaseModel):
    email: str
    password: str

class ChatRequest(BaseModel):
    email: str
    chatroom_id: int
    message: str

class QueryRequest(BaseModel):
    email: str
    query: str

class LoginRequest(BaseModel):
    email: str
    password: str

def send_email(to_email: str, subject: str, body: str) -> bool:
    if not to_email or "@" not in to_email:
        return False
    msg = MIMEMultipart()
    msg["Subject"] = subject
    msg["From"] = SMTP_USER
    msg["To"] = to_email
    msg.attach(MIMEText(body, "plain"))
    try:
        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)
        return True
    except Exception as e:
        print(f"Email failed: {e}")
        return False

def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="abhinand/MedEmbed-base-v0.1",
        model_kwargs={"device": "cpu"}
    )

def get_vectorstore():
    return Chroma(persist_directory=PERSIST_DIR, embedding_function=get_embeddings())

def contains_prescription(text: str) -> bool:
    return any(re.search(p, text, re.IGNORECASE) for p in [
        "Prescription:", "prescribed?", "take.*mg", "dosage", "medicine",
        "medication", "drug", "tablet", "capsule", "syrup", "inject",
        "twice.*day", "once.*day", "three.*times"
    ])

def refine_query(original_query: str) -> str:
    prompt = ("You are a query rewriting assistant for a medical retrieval system. "
              "Rewrite the user's query clearly. If vague, expand. If single-word, "
              "make it a medical question. Keep only the query.")
    resp = chat_llm.invoke([
        SystemMessage(content=prompt),
        HumanMessage(content=f"Original query: {original_query}")
    ])
    return resp.content.strip()

def load_docs_from_file(tmp_path: str, ext: str, source: str) -> List[Document]:
    meta = {"source": source, "processed_date": datetime.now().isoformat()}
    docs = []
    if ext == ".pdf":
        docs = [d for d in PyPDFLoader(tmp_path).load()]
        [d.metadata.update(meta) for d in docs]
    elif ext == ".docx":
        docs = [Document(page_content="\n".join(p.text for p in DocxDocument(tmp_path).paragraphs), metadata=meta)]
    elif ext == ".csv":
        docs = [d for d in CSVLoader(tmp_path, encoding="utf-8").load()]
        [d.metadata.update(meta) for d in docs]
    return splitter.split_documents(docs)

@app.post("/register")
def register(req: RegisterRequest):
    cur.execute("SELECT * FROM users WHERE email=%s", (req.email,))
    if cur.fetchone():
        raise HTTPException(status_code=400, detail="email already exists")
    
    
    hashed_password = hash_password(req.password)
    cur.execute("INSERT INTO users (email, password) VALUES (%s, %s)", (
        req.email, hashed_password))
    conn.commit()
    return {"message": "User registered successfully"}


@app.post("/login")
def login(req: LoginRequest):
    cur.execute("SELECT email, password FROM users WHERE email=%s", (req.email,))
    user = cur.fetchone()
    
   
    if not user or not verify_password(req.password, user[1]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid credentials"
        )
    
    
    access_token = create_access_token(data={"sub": req.email})
    return {
        "message": "Login successful",
        "access_token": access_token,  
        "token_type": "bearer",       
        "user": {"email": req.email}
    }


@app.get("/chatrooms/{email}")
def get_chatrooms(email: str, current_user: str = Depends(verify_token)):
    
    if current_user != email:
        raise HTTPException(status_code=403, detail="Access denied")
    
    cur.execute("SELECT id FROM users WHERE email=%s", (email,))
    user = cur.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user_id = user[0]
    cur.execute("SELECT id, created_at FROM chatrooms WHERE user_id=%s ORDER BY created_at DESC", (user_id,))
    return {"chatrooms": cur.fetchall()}


@app.get("/messages/{chatroom_id}")
def get_messages(chatroom_id: int, current_user: str = Depends(verify_token)):
   
    cur.execute("""
        SELECT c.id FROM chatrooms c 
        JOIN users u ON c.user_id = u.id 
        WHERE c.id = %s AND u.email = %s
    """, (chatroom_id, current_user))
    
    if not cur.fetchone():
        raise HTTPException(status_code=403, detail="Access denied to this chatroom")
    
    cur.execute("SELECT user_message, ai_response FROM chat_threads WHERE chatroom_id=%s", (chatroom_id,))
    return {"messages": cur.fetchall()}

@app.post("/new_chatroom/{email}")
def new_chatroom(email: str, current_user: str = Depends(verify_token)):
    
    if current_user != email:
        raise HTTPException(status_code=403, detail="Access denied")
    
    cur.execute("SELECT id FROM users WHERE email=%s", (email,))
    user = cur.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user_id = user[0]
    cur.execute("INSERT INTO chatrooms (user_id, created_at) VALUES (%s, %s) RETURNING id", (user_id, datetime.utcnow()))
    chatroom_id = cur.fetchone()[0]
    conn.commit()
    return {"chatroom_id": chatroom_id}


@app.post("/chat")
def chat(req: ChatRequest, current_user: str = Depends(verify_token)):
    
    if current_user != req.email:
        raise HTTPException(status_code=403, detail="Access denied")
    
    
    cur.execute("SELECT id, email FROM users WHERE email=%s", (req.email,))
    user = cur.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user_id, user_email = user

    
    cur.execute("SELECT id FROM chatrooms WHERE id=%s AND user_id=%s", (req.chatroom_id, user_id))
    chatroom = cur.fetchone()
    if not chatroom:
        raise HTTPException(status_code=404, detail="Chatroom not found")

    
    cur.execute("INSERT INTO chat_threads (chatroom_id, user_message) VALUES (%s, %s) RETURNING id", (req.chatroom_id, req.message))
    thread_id = cur.fetchone()[0]
    conn.commit()

    def generate():
        ai_full_response = ""
        system_msg = SystemMessage(content="You are a helpful medical assistant. If you recommend medicine, "
                                          "format as 'Prescription: <name and dosage>'. Answer ONLY based on context.")
        messages = [system_msg, HumanMessage(content=req.message)]
        for chunk in chat_llm.stream(messages):
            part = chunk.content
            ai_full_response += part
            yield part

       
        cur.execute("UPDATE chat_threads SET ai_response=%s WHERE id=%s", (ai_full_response, thread_id))
        conn.commit()

        #
        if contains_prescription(ai_full_response):
            send_email(user_email, "Prescription Alert", ai_full_response)

    return StreamingResponse(generate(), media_type="text/plain")


@app.post("/upload")
async def upload_documents(files: List[UploadFile] = File(...), current_user: str = Depends(verify_token)):
 
    if not files:
        raise HTTPException(400, "No files provided")
    vs, results = get_vectorstore(), []
    for uf in files:
        ext = os.path.splitext(uf.filename.lower())[1]
        if ext not in [".pdf", ".docx", ".csv"]:
            results.append({"filename": uf.filename, "status": "unsupported", "chunks_count": 0})
            continue
        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
                shutil.copyfileobj(uf.file, tmp)
                tmp_path = tmp.name
            chunks = load_docs_from_file(tmp_path, ext, uf.filename)
            if not chunks:
                results.append({"filename": uf.filename, "status": "no content", "chunks_count": 0})
                continue
            vs.add_documents(chunks)
            results.append({"filename": uf.filename, "status": "success", "chunks_count": len(chunks)})
        except Exception as e:
            results.append({"filename": uf.filename, "status": str(e), "chunks_count": 0})
        finally:
            tmp_path and os.path.exists(tmp_path) and os.unlink(tmp_path)
    vs.persist()
    return {"message": f"Processed {sum(r['status']=='success' for r in results)} documents", "documents": results}


@app.post("/query")
async def query_documents(req: QueryRequest, current_user: str = Depends(verify_token)):
   
    if current_user != req.email:
        raise HTTPException(status_code=403, detail="Access denied")
    
    
    cur.execute("SELECT id, email FROM users WHERE email=%s", (req.email,))
    user = cur.fetchone()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    user_id, user_email = user

    q = (req.query or "").strip()
    if len(q) < 3:
        raise HTTPException(400, "Query too short")

    vs = get_vectorstore()
    refined = refine_query(q)
    docs = vs.similarity_search(refined, k=5)
    if not docs:
        return {"answer": "No relevant info.", "sources": []}

    context = "\n\n".join(f"Doc {i+1}:\n{d.page_content}" for i, d in enumerate(docs))
    system_msg = ("You are a helpful medical assistant. "
                  "If you recommend medicine, format as 'Prescription: <name and dosage>'. "
                  "Answer based ONLY on context. If unknown, say so.")
    messages = [SystemMessage(content=system_msg), HumanMessage(content=f"Context:\n{context}\n\nQuestion: {q}")]
    answer = chat_llm.invoke(messages).content

   
    cur.execute("INSERT INTO chat_threads (chatroom_id, user_message, ai_response) VALUES "
                "((SELECT id FROM chatrooms WHERE user_id=%s ORDER BY created_at DESC LIMIT 1), %s, %s)", (user_id, q, answer))
    conn.commit()


    prescription_detected = contains_prescription(answer)
    email_sent = False
    if prescription_detected:
        email_sent = send_email(user_email, "Medical Query Response", f"Query: {q}\n\nAnswer:\n{answer}")

    sources = sorted({d.metadata.get("source", "Unknown") for d in docs})
    return {"answer": answer, "sources": sources, "email_sent": email_sent}